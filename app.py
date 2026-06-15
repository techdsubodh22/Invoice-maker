import os
import json
import io
import re
from datetime import date
from flask import Flask, render_template, request, jsonify, send_file, abort

app = Flask(__name__)

INVOICE_COUNTER_FILE = os.path.join(os.path.dirname(__file__), 'invoice_counter.json')


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _load_counter():
    if os.path.exists(INVOICE_COUNTER_FILE):
        with open(INVOICE_COUNTER_FILE, 'r') as f:
            return json.load(f)
    return {"year": date.today().year, "seq": 0}


def _save_counter(data):
    with open(INVOICE_COUNTER_FILE, 'w') as f:
        json.dump(data, f)


def _next_invoice_number():
    data = _load_counter()
    current_year = date.today().year
    if data.get("year") != current_year:
        data = {"year": current_year, "seq": 0}
    data["seq"] += 1
    _save_counter(data)
    return f"INV-{data['year']}-{data['seq']:03d}"


def _peek_invoice_number():
    """Return what the *next* number would be without incrementing."""
    data = _load_counter()
    current_year = date.today().year
    if data.get("year") != current_year:
        seq = 1
    else:
        seq = data.get("seq", 0) + 1
    return f"INV-{current_year}-{seq:03d}"


def _sanitize(data: dict) -> dict:
    """Coerce numeric fields and calculate derived totals.

    Items have an 'item_type' field: 'service' or 'expense'.
    GST applies only to service items. Expenses are reimbursements.
    """
    items = data.get("items", [])
    services_subtotal = 0.0
    expenses_subtotal = 0.0

    for item in items:
        try:
            qty  = float(item.get("quantity", 0) or 0)
            rate = float(item.get("rate", 0) or 0)
        except (ValueError, TypeError):
            qty, rate = 0.0, 0.0
        item["quantity"] = qty
        item["rate"]     = rate
        item["amount"]   = round(qty * rate, 2)
        if item.get("item_type") == "expense":
            expenses_subtotal += item["amount"]
        else:
            item["item_type"] = "service"
            services_subtotal += item["amount"]

    data["services_subtotal"] = round(services_subtotal, 2)
    data["expenses_subtotal"] = round(expenses_subtotal, 2)
    data["subtotal"]          = round(services_subtotal + expenses_subtotal, 2)
    data["has_expenses"]      = expenses_subtotal > 0

    gst_mode = data.get("gst_mode", False)

    if gst_mode:
        gst_type = data.get("gst_type", "cgst_sgst")
        # GST on services only
        if gst_type == "cgst_sgst":
            cgst = round(services_subtotal * 9 / 100, 2)
            sgst = round(services_subtotal * 9 / 100, 2)
            igst = 0.0
        else:
            cgst = 0.0
            sgst = 0.0
            igst = round(services_subtotal * 18 / 100, 2)
        data["cgst"]       = cgst
        data["sgst"]       = sgst
        data["igst"]       = igst
        data["total_gst"]  = round(cgst + sgst + igst, 2)
        data["grand_total"] = round(services_subtotal + expenses_subtotal + data["total_gst"], 2)
    else:
        data["grand_total"] = round(services_subtotal + expenses_subtotal, 2)

    return data


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    next_num = _peek_invoice_number()
    today = date.today().isoformat()
    return render_template("index.html", next_invoice_number=next_num, today=today)


@app.route("/api/next-invoice-number")
def api_next_invoice_number():
    return jsonify({"invoice_number": _peek_invoice_number()})


@app.route("/preview", methods=["POST"])
def preview():
    data = request.get_json(force=True)
    data = _sanitize(data)
    template_name = data.get("template", "classic")
    allowed = {"classic", "modern", "minimal", "corporate"}
    if template_name not in allowed:
        template_name = "classic"
    return render_template("preview.html", d=data, template=template_name)


@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    try:
        from xhtml2pdf import pisa
    except ImportError:
        abort(500, "xhtml2pdf is not installed. Run: pip install xhtml2pdf")

    data = request.get_json(force=True)
    data = _sanitize(data)
    template_name = data.get("template", "classic")
    allowed = {"classic", "modern", "minimal", "corporate"}
    if template_name not in allowed:
        template_name = "classic"

    # Bump invoice counter only on export
    if not data.get("invoice_number"):
        data["invoice_number"] = _next_invoice_number()

    html_string = render_template("preview.html", d=data, template=template_name)

    pdf_buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(html_string, dest=pdf_buffer)
    if pisa_status.err:
        abort(500, "PDF generation failed")
    pdf_bytes = pdf_buffer.getvalue()

    inv_num = re.sub(r'[^A-Za-z0-9_-]', '_', data.get("invoice_number", "invoice"))
    filename = f"{inv_num}.pdf"
    return send_file(
        io.BytesIO(pdf_bytes),
        mimetype="application/pdf",
        as_attachment=True,
        download_name=filename
    )


@app.route("/export/docx", methods=["POST"])
def export_docx():
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        abort(500, "python-docx is not installed. Run: pip install python-docx")

    data = request.get_json(force=True)
    data = _sanitize(data)

    if not data.get("invoice_number"):
        data["invoice_number"] = _next_invoice_number()

    gst_mode = data.get("gst_mode", False)
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    def add_heading(text, level=1, bold=True, size=16, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    def add_kv(label, value, bold_label=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run_l = p.add_run(f"{label}: ")
        run_l.bold = bold_label
        run_l.font.size = Pt(9)
        run_v = p.add_run(str(value) if value else "")
        run_v.font.size = Pt(9)
        return p

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    # Title
    title_p = add_heading("INVOICE", level=1, size=22, color=(30, 80, 160))
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Invoice meta
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(6)
    r = meta.add_run(f"Invoice No: {data.get('invoice_number', '')}   |   Date: {data.get('invoice_date', '')}   |   Due: {data.get('due_date', '')}")
    r.font.size = Pt(9)

    doc.add_paragraph()

    # From / To table
    ft_table = doc.add_table(rows=1, cols=2)
    ft_table.style = 'Table Grid'
    ft_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    from_cell = ft_table.rows[0].cells[0]
    to_cell = ft_table.rows[0].cells[1]

    set_cell_bg(from_cell, "EBF3FB")
    set_cell_bg(to_cell, "EBF3FB")

    def fill_party(cell, heading, fields):
        cell.paragraphs[0].clear()
        h = cell.paragraphs[0].add_run(heading)
        h.bold = True
        h.font.size = Pt(10)
        for label, val in fields:
            if val:
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                rl = p.add_run(f"{label}: ")
                rl.bold = True
                rl.font.size = Pt(8)
                rv = p.add_run(str(val))
                rv.font.size = Pt(8)

    from_fields = [
        ("Name", data.get("consultant_name")),
        ("Address", data.get("consultant_address")),
        ("Phone", data.get("consultant_phone")),
        ("Email", data.get("consultant_email")),
        ("PAN", data.get("consultant_pan")),
    ]
    if gst_mode:
        from_fields.append(("GSTIN", data.get("consultant_gstin")))

    to_fields = [
        ("Company", data.get("client_company")),
        ("Address", data.get("client_address")),
        ("Contact", data.get("client_contact")),
    ]
    if gst_mode:
        to_fields.append(("GSTIN", data.get("client_gstin")))
        to_fields.append(("Place of Supply", data.get("place_of_supply")))

    fill_party(from_cell, "FROM (Consultant)", from_fields)
    fill_party(to_cell, "TO (Client)", to_fields)

    doc.add_paragraph()

    # Line items table
    gst_cols = gst_mode
    headers = ["#", "Description", "Unit", "Qty", "Rate (₹)", "Amount (₹)"]
    if gst_cols:
        headers.insert(2, "SAC")
    col_count = len(headers)

    items_table = doc.add_table(rows=1, cols=col_count)
    items_table.style = 'Table Grid'
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = items_table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "1E50A0")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, item in enumerate(data.get("items", []), 1):
        row = items_table.add_row()
        is_expense = item.get("item_type") == "expense"
        desc = item.get("description", "")
        if is_expense:
            desc = f"[Expense] {desc}"
        vals = [str(idx), desc, item.get("unit", ""), str(item.get("quantity", "")), f"₹{item.get('rate', 0):,.2f}", f"₹{item.get('amount', 0):,.2f}"]
        if gst_cols:
            vals.insert(2, item.get("sac_code", ""))
        bg = "FFF8F0" if is_expense else ("FFFFFF" if idx % 2 == 1 else "F5F8FF")
        for i, v in enumerate(vals):
            cell = row.cells[i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(9)
            if is_expense:
                run.font.color.rgb = RGBColor(180, 90, 0)
            if i >= (3 if gst_cols else 2):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Totals
    totals_table = doc.add_table(rows=0, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    def add_total_row(label, value, bold=False, bg="FFFFFF"):
        row = totals_table.add_row()
        lc, vc = row.cells[0], row.cells[1]
        set_cell_bg(lc, bg)
        set_cell_bg(vc, bg)
        rl = lc.paragraphs[0].add_run(label)
        rl.bold = bold
        rl.font.size = Pt(9)
        rv = vc.paragraphs[0].add_run(value)
        rv.bold = bold
        rv.font.size = Pt(9)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_total_row("Services Subtotal", f"₹{data['services_subtotal']:,.2f}")
    if data.get("has_expenses"):
        add_total_row("Reimbursable Expenses", f"₹{data['expenses_subtotal']:,.2f}")
    if gst_mode:
        gst_type = data.get("gst_type", "cgst_sgst")
        if gst_type == "cgst_sgst":
            add_total_row("CGST (9%) on Services", f"₹{data['cgst']:,.2f}")
            add_total_row("SGST (9%) on Services", f"₹{data['sgst']:,.2f}")
        else:
            add_total_row("IGST (18%) on Services", f"₹{data['igst']:,.2f}")
        add_total_row("Total GST", f"₹{data['total_gst']:,.2f}")
    add_total_row("GRAND TOTAL", f"₹{data['grand_total']:,.2f}", bold=True, bg="1E50A0")

    # Fix grand total text color
    last_row = totals_table.rows[-1]
    for cell in last_row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    if data.get("show_bank"):
        doc.add_paragraph()
        bank_heading = doc.add_paragraph()
        bh = bank_heading.add_run("Payment Details")
        bh.bold = True
        bh.font.size = Pt(10)
        bank_fields = [
            ("Bank Name", data.get("bank_name")),
            ("Account Number", data.get("account_number")),
            ("IFSC Code", data.get("ifsc_code")),
            ("Account Type", data.get("account_type")),
        ]
        for label, val in bank_fields:
            if val:
                add_kv(label, val)

    if data.get("show_notes") and data.get("notes"):
        doc.add_paragraph()
        nh = doc.add_paragraph()
        nh.add_run("Notes / Terms").bold = True
        notes_p = doc.add_paragraph(data.get("notes", ""))
        if notes_p.runs:
            notes_p.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    inv_num = re.sub(r'[^A-Za-z0-9_-]', '_', data.get("invoice_number", "invoice"))
    filename = f"{inv_num}.docx"
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )


if __name__ == "__main__":
    app.run(debug=True, port=7765)
